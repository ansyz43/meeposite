"""Content plan router — /api/content

Isolated from all existing routers. No shared state.
"""
import io
import logging

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select, func as sa_func, delete
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.auth import get_current_user
from app.database import get_db
from app.models import (
    CompetitorPost,
    CompetitorSource,
    ContentPlan,
    ContentPlanItem,
    ContentProfile,
    User,
)
from app.schemas import (
    CompetitorPostResponse,
    CompetitorSourceAdd,
    CompetitorSourceResponse,
    ContentPlanItemUpdate,
    ContentPlanListItem,
    ContentPlanResponse,
    ContentProfileRequest,
    ContentProfileResponse,
    GeneratePlanRequest,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/content", tags=["content"])


# ── Profile ──────────────────────────────────────────────────

@router.get("/profile", response_model=ContentProfileResponse | None)
async def get_profile(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(ContentProfile)
        .options(selectinload(ContentProfile.competitor_sources))
        .where(ContentProfile.user_id == user.id)
    )
    profile = result.scalar_one_or_none()
    if not profile:
        return None

    # Build response with competitor post counts
    competitors = []
    for src in profile.competitor_sources:
        count = await db.execute(
            select(sa_func.count(CompetitorPost.id)).where(
                CompetitorPost.platform == src.platform,
                CompetitorPost.channel_username == src.channel_username,
            )
        )
        competitors.append(CompetitorSourceResponse(
            id=src.id,
            platform=src.platform,
            channel_username=src.channel_username,
            channel_title=src.channel_title,
            last_parsed_at=src.last_parsed_at,
            is_active=src.is_active,
            post_count=count.scalar() or 0,
        ))

    return ContentProfileResponse(
        id=profile.id,
        niche=profile.niche,
        platforms=profile.platforms.split(",") if profile.platforms else [],
        tone=profile.tone,
        target_audience=profile.target_audience,
        topics=profile.topics.split(",") if profile.topics else [],
        founder_story=profile.founder_story,
        transformation=profile.transformation,
        meepo_bot_deeplink=profile.meepo_bot_deeplink,
        created_at=profile.created_at,
        updated_at=profile.updated_at,
        competitors=competitors,
    )


@router.post("/profile", response_model=ContentProfileResponse)
async def upsert_profile(
    data: ContentProfileRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(ContentProfile).where(ContentProfile.user_id == user.id)
    )
    profile = result.scalar_one_or_none()

    platforms_str = ",".join(data.platforms)
    topics_str = ",".join(data.topics)

    # Detect changes that should invalidate the cached narrative core.
    core_inputs_changed = False

    if profile:
        if (
            profile.niche != data.niche
            or profile.tone != data.tone
            or (profile.target_audience or "") != (data.target_audience or "")
            or (profile.topics or "") != topics_str
            or (profile.founder_story or "") != (data.founder_story or "")
            or (profile.transformation or "") != (data.transformation or "")
        ):
            core_inputs_changed = True

        profile.niche = data.niche
        profile.platforms = platforms_str
        profile.tone = data.tone
        profile.target_audience = data.target_audience
        profile.topics = topics_str
        profile.founder_story = data.founder_story or None
        profile.transformation = data.transformation or None
        profile.meepo_bot_deeplink = data.meepo_bot_deeplink or None

        if core_inputs_changed:
            # Drop cached strategy — it will regenerate on the next plan.
            profile.strategy_json = None
            profile.strategy_generated_at = None
    else:
        profile = ContentProfile(
            user_id=user.id,
            niche=data.niche,
            platforms=platforms_str,
            tone=data.tone,
            target_audience=data.target_audience,
            topics=topics_str,
            founder_story=data.founder_story or None,
            transformation=data.transformation or None,
            meepo_bot_deeplink=data.meepo_bot_deeplink or None,
        )
        db.add(profile)

    await db.commit()
    await db.refresh(profile)

    return ContentProfileResponse(
        id=profile.id,
        niche=profile.niche,
        platforms=data.platforms,
        tone=profile.tone,
        target_audience=profile.target_audience,
        topics=data.topics,
        founder_story=profile.founder_story,
        transformation=profile.transformation,
        meepo_bot_deeplink=profile.meepo_bot_deeplink,
        created_at=profile.created_at,
        updated_at=profile.updated_at,
        competitors=[],
    )


# ── Competitors ──────────────────────────────────────────────

@router.post("/competitors", response_model=CompetitorSourceResponse)
async def add_competitor(
    data: CompetitorSourceAdd,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Must have profile first
    result = await db.execute(
        select(ContentProfile).where(ContentProfile.user_id == user.id)
    )
    profile = result.scalar_one_or_none()
    if not profile:
        raise HTTPException(400, "Сначала создайте профиль")

    # Check limit (max 10 competitors)
    count = await db.execute(
        select(sa_func.count(CompetitorSource.id)).where(
            CompetitorSource.profile_id == profile.id,
        )
    )
    if (count.scalar() or 0) >= 10:
        raise HTTPException(400, "Максимум 10 конкурентов")

    # Check duplicate
    existing = await db.execute(
        select(CompetitorSource).where(
            CompetitorSource.profile_id == profile.id,
            CompetitorSource.platform == data.platform,
            CompetitorSource.channel_username == data.channel_username,
        )
    )
    if existing.scalar_one_or_none():
        raise HTTPException(400, "Этот конкурент уже добавлен")

    source = CompetitorSource(
        profile_id=profile.id,
        platform=data.platform,
        channel_username=data.channel_username,
    )
    db.add(source)
    await db.commit()
    await db.refresh(source)

    # Trigger parse in background
    background_tasks.add_task(_parse_competitor_bg, data.platform, data.channel_username)

    return CompetitorSourceResponse(
        id=source.id,
        platform=source.platform,
        channel_username=source.channel_username,
        channel_title=source.channel_title,
        last_parsed_at=source.last_parsed_at,
        is_active=source.is_active,
        post_count=0,
    )


@router.delete("/competitors/{source_id}")
async def delete_competitor(
    source_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(CompetitorSource)
        .join(ContentProfile)
        .where(
            CompetitorSource.id == source_id,
            ContentProfile.user_id == user.id,
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(404, "Конкурент не найден")

    await db.delete(source)
    await db.commit()
    return {"ok": True}


@router.get("/competitors/{source_id}/posts", response_model=list[CompetitorPostResponse])
async def get_competitor_posts(
    source_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(CompetitorSource)
        .join(ContentProfile)
        .where(
            CompetitorSource.id == source_id,
            ContentProfile.user_id == user.id,
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(404, "Конкурент не найден")

    posts = await db.execute(
        select(CompetitorPost)
        .where(
            CompetitorPost.platform == source.platform,
            CompetitorPost.channel_username == source.channel_username,
        )
        .order_by(CompetitorPost.posted_at.desc())
        .limit(30)
    )
    return [
        CompetitorPostResponse(
            id=p.id,
            text=p.text[:500],  # truncate for listing
            views=p.views,
            reactions=p.reactions,
            posted_at=p.posted_at,
        )
        for p in posts.scalars().all()
    ]


@router.post("/competitors/{source_id}/parse")
async def parse_competitor(
    source_id: int,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Manually trigger re-parse of a competitor."""
    result = await db.execute(
        select(CompetitorSource)
        .join(ContentProfile)
        .where(
            CompetitorSource.id == source_id,
            ContentProfile.user_id == user.id,
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(404, "Конкурент не найден")

    background_tasks.add_task(_parse_competitor_bg, source.platform, source.channel_username, force=True)
    return {"ok": True, "message": "Парсинг запущен"}


# ── Plans ────────────────────────────────────────────────────

@router.get("/plans", response_model=list[ContentPlanListItem])
async def list_plans(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(ContentPlan)
        .where(ContentPlan.user_id == user.id)
        .order_by(ContentPlan.created_at.desc())
        .limit(20)
    )
    plans = result.scalars().all()

    items = []
    for plan in plans:
        count = await db.execute(
            select(sa_func.count(ContentPlanItem.id)).where(ContentPlanItem.plan_id == plan.id)
        )
        items.append(ContentPlanListItem(
            id=plan.id,
            title=plan.title,
            platform=plan.platform,
            period_days=plan.period_days,
            status=plan.status,
            created_at=plan.created_at,
            item_count=count.scalar() or 0,
        ))

    return items


@router.get("/plans/{plan_id}", response_model=ContentPlanResponse)
async def get_plan(
    plan_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(ContentPlan)
        .options(selectinload(ContentPlan.items))
        .where(ContentPlan.id == plan_id, ContentPlan.user_id == user.id)
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "План не найден")

    return ContentPlanResponse(
        id=plan.id,
        title=plan.title,
        platform=plan.platform,
        period_days=plan.period_days,
        status=plan.status,
        error_message=plan.error_message,
        created_at=plan.created_at,
        items=[
            {
                "id": item.id,
                "day_number": item.day_number,
                "post_type": item.post_type,
                "topic": item.topic,
                "text": item.text,
                "hashtags": item.hashtags,
                "best_time": item.best_time,
                "script": item.script,
                "hunt_stage": item.hunt_stage,
                "is_meepo_cta": item.is_meepo_cta,
                "is_edited": item.is_edited,
            }
            for item in plan.items
        ],
    )


@router.post("/plans/generate", response_model=ContentPlanResponse, status_code=202)
async def generate_plan(
    data: GeneratePlanRequest,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Create a plan row with status='generating' and run generation in background.

    The frontend should poll GET /plans/{id} until status becomes 'ready' or 'error'.
    """
    from app.services.content_ai import fill_plan_background

    # Make sure the user has a profile (fail fast before creating a placeholder plan).
    profile_res = await db.execute(
        select(ContentProfile).where(ContentProfile.user_id == user.id)
    )
    if not profile_res.scalar_one_or_none():
        raise HTTPException(400, "Сначала заполните профиль контент-маркетинга")

    platform_name = "Instagram" if data.platform == "instagram" else "Telegram"
    plan = ContentPlan(
        user_id=user.id,
        title=f"Контент-план {platform_name} на {data.period_days} дней",
        platform=data.platform,
        period_days=data.period_days,
        status="generating",
    )
    db.add(plan)
    await db.commit()
    await db.refresh(plan)

    background_tasks.add_task(fill_plan_background, plan.id)

    return ContentPlanResponse(
        id=plan.id,
        title=plan.title,
        platform=plan.platform,
        period_days=plan.period_days,
        status=plan.status,
        error_message=plan.error_message,
        created_at=plan.created_at,
        items=[],
    )


@router.put("/plans/{plan_id}/items/{item_id}", response_model=dict)
async def update_plan_item(
    plan_id: int,
    item_id: int,
    data: ContentPlanItemUpdate,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Verify ownership
    result = await db.execute(
        select(ContentPlanItem)
        .join(ContentPlan)
        .where(
            ContentPlanItem.id == item_id,
            ContentPlanItem.plan_id == plan_id,
            ContentPlan.user_id == user.id,
        )
    )
    item = result.scalar_one_or_none()
    if not item:
        raise HTTPException(404, "Пост не найден")

    if data.text is not None:
        item.text = data.text
    if data.hashtags is not None:
        item.hashtags = data.hashtags
    if data.topic is not None:
        item.topic = data.topic
    if data.script is not None:
        item.script = data.script
    item.is_edited = True

    await db.commit()
    return {"ok": True}


@router.delete("/plans/{plan_id}")
async def delete_plan(
    plan_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(ContentPlan).where(ContentPlan.id == plan_id, ContentPlan.user_id == user.id)
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "План не найден")

    await db.delete(plan)
    await db.commit()
    return {"ok": True}


# ── Export ────────────────────────────────────────────────────

async def _load_plan_for_export(plan_id: int, user: User, db: AsyncSession) -> ContentPlan:
    result = await db.execute(
        select(ContentPlan)
        .options(selectinload(ContentPlan.items))
        .where(ContentPlan.id == plan_id, ContentPlan.user_id == user.id)
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "План не найден")
    if plan.status != "ready":
        raise HTTPException(400, "План ещё не готов")
    return plan


@router.get("/plans/{plan_id}/export/pdf")
async def export_plan_pdf(
    plan_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    plan = await _load_plan_for_export(plan_id, user, db)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=20*mm, rightMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)

    # Try to register a font that supports Cyrillic
    font_name = "Helvetica"
    font_name_bold = "Helvetica-Bold"
    for font_path in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("CyrFont", font_path))
            bold_path = font_path.replace("Sans.ttf", "Sans-Bold.ttf").replace("arial.ttf", "arialbd.ttf")
            if os.path.exists(bold_path):
                pdfmetrics.registerFont(TTFont("CyrFontBold", bold_path))
                font_name_bold = "CyrFontBold"
            font_name = "CyrFont"
            break

    # Styles
    accent = HexColor("#10b981")
    dark = HexColor("#1a1a2e")
    gray = HexColor("#6b7280")

    s_title = ParagraphStyle("Title", fontName=font_name_bold, fontSize=18, textColor=dark, alignment=TA_CENTER, spaceAfter=4*mm)
    s_sub = ParagraphStyle("Sub", fontName=font_name, fontSize=10, textColor=gray, alignment=TA_CENTER, spaceAfter=8*mm)
    s_day = ParagraphStyle("Day", fontName=font_name_bold, fontSize=12, textColor=accent, spaceBefore=6*mm, spaceAfter=2*mm)
    s_label = ParagraphStyle("Label", fontName=font_name_bold, fontSize=9, textColor=gray, spaceBefore=2*mm, spaceAfter=1*mm)
    s_body = ParagraphStyle("Body", fontName=font_name, fontSize=10, textColor=dark, leading=14, spaceAfter=2*mm)
    s_meta = ParagraphStyle("Meta", fontName=font_name, fontSize=8, textColor=gray, spaceAfter=1*mm)
    s_script = ParagraphStyle("Script", fontName=font_name, fontSize=10, textColor=HexColor("#7c3aed"), leading=14, leftIndent=6*mm, spaceAfter=2*mm)
    s_hash = ParagraphStyle("Hash", fontName=font_name, fontSize=8, textColor=accent, spaceAfter=4*mm)

    story = []
    story.append(Paragraph(plan.title.replace("&", "&amp;"), s_title))
    platform_label = "Instagram" if plan.platform == "instagram" else "Telegram"
    story.append(Paragraph(f"{platform_label} · {plan.period_days} дней · {plan.created_at.strftime('%d.%m.%Y')}", s_sub))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#e5e7eb"), spaceAfter=4*mm))

    type_emoji = {"пост": "📝", "сторис": "📱", "рилс": "🎬", "карусель": "🖼"}

    for item in plan.items:
        emoji = type_emoji.get(item.post_type, "📌")
        story.append(Paragraph(f"День {item.day_number} — {emoji} {item.post_type.upper()}", s_day))
        if item.best_time:
            story.append(Paragraph(f"⏰ Лучшее время: {item.best_time}", s_meta))
        story.append(Paragraph(f"Тема: {item.topic.replace('&', '&amp;')}", s_label))
        # Escape XML special chars for reportlab
        text_safe = item.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        story.append(Paragraph(text_safe, s_body))

        if item.script:
            story.append(Paragraph("🎤 Скрипт для озвучки:", s_label))
            script_safe = item.script.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            story.append(Paragraph(script_safe, s_script))

        if item.hashtags:
            story.append(Paragraph(item.hashtags.replace("&", "&amp;"), s_hash))

        story.append(HRFlowable(width="100%", thickness=0.3, color=HexColor("#e5e7eb"), spaceAfter=2*mm))

    doc.build(story)
    buf.seek(0)

    filename = f"content_plan_{plan.id}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/plans/{plan_id}/export/docx")
async def export_plan_docx(
    plan_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    plan = await _load_plan_for_export(plan_id, user, db)

    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    document = Document()

    # Narrow margins
    for section in document.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title
    title_p = document.add_heading(plan.title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    # Subtitle
    platform_label = "Instagram" if plan.platform == "instagram" else "Telegram"
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{platform_label} · {plan.period_days} дней · {plan.created_at.strftime('%d.%m.%Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    document.add_paragraph()  # spacer

    type_emoji = {"пост": "📝", "сторис": "📱", "рилс": "🎬", "карусель": "🖼"}

    for item in plan.items:
        emoji = type_emoji.get(item.post_type, "📌")

        # Day heading
        h = document.add_heading(f"День {item.day_number} — {emoji} {item.post_type.upper()}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
            run.font.size = Pt(13)

        # Meta
        if item.best_time:
            meta_p = document.add_paragraph()
            run = meta_p.add_run(f"⏰ Лучшее время: {item.best_time}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Topic
        topic_p = document.add_paragraph()
        run = topic_p.add_run("Тема: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run = topic_p.add_run(item.topic)
        run.font.size = Pt(10)

        # Post text
        text_p = document.add_paragraph()
        run = text_p.add_run(item.text)
        run.font.size = Pt(10)
        text_p.paragraph_format.space_after = Pt(4)

        # Script
        if item.script:
            script_label = document.add_paragraph()
            run = script_label.add_run("🎤 Скрипт для озвучки:")
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

            script_p = document.add_paragraph()
            script_p.paragraph_format.left_indent = Cm(1)
            run = script_p.add_run(item.script)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            run.font.italic = True
            script_p.paragraph_format.space_after = Pt(4)

        # Hashtags
        if item.hashtags:
            hash_p = document.add_paragraph()
            run = hash_p.add_run(item.hashtags)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

        # Separator
        document.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    filename = f"content_plan_{plan.id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Auto-detect profile ──────────────────────────────────────

@router.post("/profile/auto-detect")
async def auto_detect_profile(
    data: dict,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Parse user's own social media and auto-detect niche, audience, tone, topics."""
    platform = data.get("platform", "").lower()
    username = data.get("username", "").strip().lstrip("@").lower()

    # Extract username from Instagram URL if full link was pasted
    if platform == "instagram":
        from app.services.parser_instagram import _extract_username
        username = _extract_username(username)

    if not username:
        raise HTTPException(400, "Укажите username аккаунта")
    if platform not in ("telegram", "instagram"):
        raise HTTPException(400, "Поддерживается только Telegram и Instagram")

    # Parse posts using existing parsers
    try:
        if platform == "telegram":
            from app.services.parser_telegram import parse_telegram_channel
            result = await parse_telegram_channel(username, db, force=True, max_posts=30)
        else:
            from app.services.parser_instagram import parse_instagram_profile
            result = await parse_instagram_profile(username, db, force=True)

        if result.get("status") == "error":
            raise HTTPException(400, result.get("error", "Ошибка парсинга"))
        if result.get("status") == "rate_limited":
            raise HTTPException(429, "Слишком много запросов, попробуйте через минуту")
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Auto-detect parse failed: %s", e, exc_info=True)
        raise HTTPException(400, f"Не удалось загрузить посты: {e}")

    # Fetch parsed posts from DB
    posts_result = await db.execute(
        select(CompetitorPost.text)
        .where(
            CompetitorPost.platform == platform,
            CompetitorPost.channel_username == username,
        )
        .order_by(CompetitorPost.posted_at.desc())
        .limit(20)
    )
    posts_texts = [row[0] for row in posts_result.all() if row[0]]

    if not posts_texts:
        raise HTTPException(400, "Не найдено постов для анализа")

    # AI analysis
    from app.services.content_ai import analyze_profile_from_posts

    try:
        profile_data = await analyze_profile_from_posts(posts_texts)
    except Exception as e:
        logger.error("Auto-detect AI failed: %s", e, exc_info=True)
        raise HTTPException(500, "Ошибка AI-анализа, попробуйте ещё раз")

    return {
        "niche": profile_data.get("niche", ""),
        "target_audience": profile_data.get("target_audience", ""),
        "tone": profile_data.get("tone", "friendly"),
        "topics": profile_data.get("topics", []),
    }


# ── Background tasks ─────────────────────────────────────────

async def _parse_competitor_bg(platform: str, username: str, force: bool = False):
    """Background task: parse a competitor channel/profile."""
    from app.database import async_session

    async with async_session() as db:
        try:
            if platform == "telegram":
                from app.services.parser_telegram import parse_telegram_channel
                result = await parse_telegram_channel(username, db, force=force)
            elif platform == "instagram":
                from app.services.parser_instagram import parse_instagram_profile
                result = await parse_instagram_profile(username, db, force=force)
            else:
                logger.warning("Unknown platform: %s", platform)
                return

            logger.info("Background parse %s @%s: %s", platform, username, result.get("status"))
        except Exception as e:
            logger.error("Background parse error %s @%s: %s", platform, username, e, exc_info=True)
